"""
Stroke Analysis Data Processing Module

This module handles the processing and analysis of stroke patient data from multiple sources:
- Actigraph data (movement/activity measurements)
- ICH (Intracerebral Hemorrhage) data
- Word document medical records

The module provides functions for data preprocessing, feature extraction, and analysis.
"""

import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from scipy.spatial.distance import euclidean
from fastdtw import fastdtw
from docx import Document
from dateutil.relativedelta import relativedelta
import tensorflow as tf
from tqdm import tqdm
import os
import io
import csv
from tsfresh.feature_extraction import (extract_features, MinimalFCParameters,
                                        feature_calculators)
from deep_learning.s2s_fa import get_pretrained_encoder

# Global number of patients
NUM_PATIENTS = 41

# Global reference data
REFERENCE_DAY = pd.DataFrame(index=range(1, NUM_PATIENTS), columns=['day_num', 'date'])

#------------------------------------------------------------------------------
# Control Day Functions
#------------------------------------------------------------------------------

def get_control_day(patient_id):
    """
    Get the control (non-delirious) day for a patient.
    
    Args:
        patient_id: Patient identifier
    
    Returns:
        date: Control day date for the patient
    
    Note: Some patients (1,18,19,20,25,26,27,28,29,30,33,34,35,37,38) did not have non-delirious days
    """
    date = REFERENCE_DAY['date'][patient_id]
    return date

def populate_reference_dates(patient_id, min_date):
    """
    Convert patient's reference day number into a date based on their minimum date.
    
    Args:
        patient_id: Patient identifier
        min_date: Minimum date for the patient
    """
    ref_day_num = REFERENCE_DAY['day_num'][patient_id]-1
    ref_date_relative = pd.to_timedelta(ref_day_num, unit='D')
    ref_date = ref_date_relative+min_date
    REFERENCE_DAY['date'][patient_id] = ref_date

def populate_REFERENCE_DAYs(ich_patient_clean):
    """
    Get reference day numbers for all patients from ICH data.
    
    Args:
        ich_patient_clean: Cleaned ICH patient data
    """
    for patient_id in range(1, NUM_PATIENTS):
        patient_days = ich_patient_clean.loc[ich_patient_clean['Patient_ID'] == patient_id]
        patient_days = patient_days.loc[patient_days['delirium_stat'] == 0]
        patient_days = patient_days.sort_values(by=['Day_Num'])
        patient_days = patient_days[patient_days['Day_Num'] != 1]
        
        ref_day_num = 2 if patient_days.empty else patient_days.reset_index()['Day_Num'][0]
        REFERENCE_DAY['day_num'][patient_id] = ref_day_num

#------------------------------------------------------------------------------
# Time Cut Functions
#------------------------------------------------------------------------------

def get_1_to_1_cuts(date):
    """Get time window from 1PM previous day to 1PM current day."""
    start_cut = date-relativedelta(days=1)+relativedelta(hours=13)
    end_cut = date+relativedelta(hours=13)
    return start_cut, end_cut

def get_night_cuts(date):
    """Get time window from 10PM previous day to 6AM current day."""
    start_cut = date-relativedelta(days=1)+relativedelta(hours=22)
    end_cut = date+relativedelta(hours=6)
    return start_cut, end_cut

def get_morning_cuts(date):
    """Get time window from 6AM to 1PM current day."""
    start_cut = date+relativedelta(hours=6)
    end_cut = date+relativedelta(hours=13)
    return start_cut, end_cut

def get_afternoon_cuts(date):
    """Get time window from 2PM to 10PM previous day."""
    start_cut = date-relativedelta(days=1)+relativedelta(hours=14)
    end_cut = date-relativedelta(days=1)+relativedelta(hours=22)
    return start_cut, end_cut

#------------------------------------------------------------------------------
# Feature Construction
#------------------------------------------------------------------------------

def construct_feature(feature_constructor, part_of_day, ACTIGRAPH_FILES,
                     patient_id, date_i, measurement=None, side=None):
    """
    Master function for adding actigraph features.
    
    Args:
        feature_constructor: Function that processes df_cut and returns value
        part_of_day: Time window to analyze
        ACTIGRAPH_FILES: Dictionary of actigraph data
        patient_id: Patient identifier
        date_i: Date to analyze
        measurement: Optional specific measurement to analyze
        side: Optional body side to analyze
    
    Returns:
        pandas.Series: Constructed feature
    """
    output_series = pd.Series()
    function_name_cut = 'get_'+part_of_day+'_cuts'
    function_cut = eval(function_name_cut)
    start_cut, end_cut = function_cut(date_i)
    
    df_cut = ACTIGRAPH_FILES[patient_id].loc[
        (ACTIGRAPH_FILES[patient_id].full_date >= start_cut) & 
        (ACTIGRAPH_FILES[patient_id].full_date <= end_cut)
    ]
    
    if df_cut.shape[0] > 0:
        if measurement is not None and side is not None:
            column_name = measurement + "_" + side
            feature_name, feature_value = feature_constructor(df_cut, part_of_day, column_name)
        else:
            feature_name, feature_value = feature_constructor(df_cut, part_of_day)
        output_series[feature_name] = feature_value
    return output_series

def construct_feature_two_day(feature_constructor, part_of_day, ACTIGRAPH_FILES,
                            patient_id, date_i, second_date, measurement=None, side=None):
    """
    Master function for adding actigraph features comparing two days.
    Similar to construct_feature but handles two dates for comparison.
    """
    output_series = pd.Series()
    function_name_cut = 'get_'+part_of_day+'_cuts'
    function_cut = eval(function_name_cut)
    
    start_cut, end_cut = function_cut(date_i)
    second_start_cut, second_end_cut = function_cut(second_date)
    
    df_cut_1 = ACTIGRAPH_FILES[patient_id].loc[
        (ACTIGRAPH_FILES[patient_id].full_date >= start_cut) & 
        (ACTIGRAPH_FILES[patient_id].full_date <= end_cut)
    ]
    
    df_cut_2 = ACTIGRAPH_FILES[patient_id].loc[
        (ACTIGRAPH_FILES[patient_id].full_date >= second_start_cut) & 
        (ACTIGRAPH_FILES[patient_id].full_date <= second_end_cut)
    ]
    
    if df_cut_1.shape[0] > 0:
        if measurement is not None and side is not None:
            column_name = measurement + "_" + side
            feature_name, feature_value = feature_constructor(df_cut_1, df_cut_2, part_of_day, column_name)
        else:
            feature_name, feature_value = feature_constructor(df_cut_1, df_cut_2, part_of_day)
        output_series[feature_name] = feature_value
    return output_series

#------------------------------------------------------------------------------
# Feature Engineering
#------------------------------------------------------------------------------

def min_at_rest(df_cut, part_of_day):
    """Calculate percentage of time at rest."""
    value = df_cut.loc[(df_cut.pim_aff == 0) & (df_cut.pim_non_aff == 0)].shape[0] / df_cut.shape[0]
    name = 'minutes_at_rest_' + part_of_day
    return name, value

def dynamic_time_warping_arms(df_cut, part_of_day):
    """Calculate DTW distance between affected and non-affected arms."""
    x = df_cut.pim_aff[~np.isnan(df_cut.pim_aff)]
    y = df_cut.pim_non_aff[~np.isnan(df_cut.pim_non_aff)]
    distance, _ = fastdtw(x, y, dist=euclidean)
    name = 'dynamic_time_warping_pim_' + part_of_day
    return name, distance

def dynamic_time_warping_two_day(df_cut_1, df_cut_2, part_of_day):
    """Calculate DTW distance between same arm on two different days."""
    x = df_cut_1.pim_aff[~np.isnan(df_cut_1.pim_aff)]
    y = df_cut_2.pim_aff[~np.isnan(df_cut_2.pim_aff)]
    distance, _ = fastdtw(x, y, dist=euclidean)
    name = 'dynamic_time_warping_two_day_pim_' + part_of_day
    return name, distance

def num_peaks(df_cut, part_of_day):
    """Calculate number of activity peaks."""
    value = feature_calculators.number_peaks(df_cut.pim_aff, 60)
    name = 'n_peaks_' + part_of_day
    return name, value

def autoencode(ACTIGRAPH_FILES, patient_id, date_i):
    """Extract latent features using pretrained autoencoder."""
    autoencoder, num_units = get_pretrained_encoder()
    output_series = pd.Series()
    start_cut, end_cut = get_1_to_1_cuts(date_i)
    
    df_cut = ACTIGRAPH_FILES[patient_id].loc[
        (ACTIGRAPH_FILES[patient_id].full_date >= start_cut) & 
        (ACTIGRAPH_FILES[patient_id].full_date <= end_cut)
    ]
    
    if df_cut.shape[0] > 0:
        encoder_input = df_cut[['pim_aff', 'pim_non_aff', "hpim_aff",
                               "hpim_non_aff", "zcm_aff", "zcm_non_aff"]].to_numpy()
        encoder_input = tf.reshape(encoder_input, (1, -1, 6))
        hidden = [tf.zeros((1, num_units))]
        _, ae_hidden = autoencoder(encoder_input, hidden)
        value = ae_hidden.numpy().flatten()
        
        for i, val in enumerate(value):
            output_series[f"ae{i+1}"] = val
            
    return output_series

#------------------------------------------------------------------------------
# Statistical Feature Functions
#------------------------------------------------------------------------------

def median_act(df_cut, part_of_day, column_name):
    """Calculate median activity value."""
    movement = df_cut.loc[df_cut[column_name] > 0]
    value = movement[column_name].median()
    name = f'median_{column_name}_{part_of_day}'
    return name, value

def max_act(df_cut, part_of_day, column_name):
    """Calculate maximum activity value."""
    movement = df_cut.loc[df_cut[column_name] > 0]
    value = movement[column_name].max()
    name = f'max_{column_name}_{part_of_day}'
    return name, value

def min_act(df_cut, part_of_day, column_name):
    """Calculate minimum activity value."""
    movement = df_cut.loc[df_cut[column_name] > 0]
    value = movement[column_name].min()
    name = f'min_{column_name}_{part_of_day}'
    return name, value

def std_act(df_cut, part_of_day, column_name):
    """Calculate standard deviation of activity."""
    movement = df_cut.loc[df_cut[column_name] > 0]
    value = movement[column_name].std()
    name = f'std_{column_name}_{part_of_day}'
    return name, value

#------------------------------------------------------------------------------
# Data Loading and Processing Functions 
#------------------------------------------------------------------------------

def load_actigraph_files(data_path):
    """Load and preprocess actigraph CSV files."""
    MERGED_FILES = {}
    
    # Process files 1-9
    for i in range(1, 10):
        filename = f'DW-00{i}-both_merged.csv'
        MERGED_FILES[i] = pd.read_csv(os.path.join(data_path, filename))
        MERGED_FILES[i]['full_date'] = MERGED_FILES[i]['date'] + ' ' + MERGED_FILES[i]['time']
        MERGED_FILES[i].date = pd.to_datetime(MERGED_FILES[i].date)
        MERGED_FILES[i].full_date = pd.to_datetime(MERGED_FILES[i].full_date)
    
    # Process files 10-40
    for i in range(10, NUM_PATIENTS):
        filename = f'DW-0{i}-both_merged.csv'
        MERGED_FILES[i] = pd.read_csv(os.path.join(data_path, filename))
        MERGED_FILES[i]['full_date'] = MERGED_FILES[i]['date'] + ' ' + MERGED_FILES[i]['time']
        MERGED_FILES[i].date = pd.to_datetime(MERGED_FILES[i].date)
        MERGED_FILES[i].full_date = pd.to_datetime(MERGED_FILES[i].full_date)
    
    return MERGED_FILES

def rename_columns_actigraph_files(MERGED_FILES, patient_location):
    """Rename columns to use affected/non-affected terminology."""
    for i in MERGED_FILES:
        affected_location = patient_location.loc[i].location_body
        non_affected_location = 'left' if affected_location == 'right' else 'right'
        
        MERGED_FILES[i].columns = MERGED_FILES[i].columns.str.replace(affected_location, "aff")
        MERGED_FILES[i].columns = MERGED_FILES[i].columns.str.replace(non_affected_location, "non_aff")
    
    return MERGED_FILES

def get_day_patient(MERGED_FILES):
    """Add day_in_hospital column to actigraph data."""
    for i in MERGED_FILES:
        MERGED_FILES[i] = MERGED_FILES[i].loc[MERGED_FILES[i].date.notna()]
        min_date = MERGED_FILES[i].date.min()
        MERGED_FILES[i]['day_in_hospital'] = (MERGED_FILES[i].date - min_date).dt.days + 1
        MERGED_FILES[i]['patient_id'] = i
    return MERGED_FILES

def remove_outliers(MERGED_FILES):
    """Removes any rows where hpim is 10 standard deviations above the mean.
    
    Args:
        MERGED_FILES (dict): Dictionary of dataframes to clean
        
    Returns:
        dict: Cleaned dataframes with outliers removed
    """
    for i in MERGED_FILES:
        df = MERGED_FILES[i]
        df = df[60:]  # Remove first 60 rows
        df = df[:-60]  # Remove last 60 rows
        std = df.std(axis=0, skipna=True)
        mean = df.mean(axis=0, skipna=True)
        df = df[df["hpim_aff"] < (mean["hpim_aff"] + 10*std["hpim_aff"])]
        MERGED_FILES[i] = df
    return MERGED_FILES

def select_features(df, use_actigraph, use_ich):
    """Select relevant features based on analysis type.
    
    Args:
        df (pd.DataFrame): Input dataframe
        use_actigraph (bool): Whether to use actigraph data
        use_ich (bool): Whether to use ICH data
        
    Returns:
        pd.DataFrame: Dataframe with selected features
    """
    df = df.dropna()
    clinical_features = [
        "age", "ich_score", "nihss_total", "ivh", "ich_side", "sex", 
        "ich_max_volume", "nihss_arm_affected", "nihss_arm_not_affected"
    ]
    
    if use_actigraph and not use_ich:
        return df.drop(clinical_features, axis=1)
    if use_ich and not use_actigraph:
        clinical_features.extend(['Patient_ID', 'Day_Num', 'delirium_stat'])
        return df[clinical_features]
    assert use_actigraph and use_ich, "error: must specify one or both kinds of data (actigraph and/or ich)"
    return df

# =====================
# Word Document Processing
# =====================

def read_docx_tables(filename, tab_id=None, **kwargs):
    """Parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)

    Args:
        filename (str): File name of a Word Document
        tab_id (int, optional): Parse a single table with the index (counting from 0)
        **kwargs: Arguments to pass to pd.read_csv()

    Returns:
        pd.DataFrame or list: Single DataFrame if tab_id specified, otherwise list of DataFrames
    """
    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)

    doc = Document(filename)
    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
    else:
        try:
            return read_docx_tab(doc.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise

def read_dw_word_files(filename):
    """Read data from tables in DWx.docx files.
    
    Args:
        filename (str): DWx.docx filename
        
    Returns:
        dict: Dictionary of pandas DataFrames, key is page number
    """
    tab = read_docx_tables(filename, None, header=None)
    output_dict = {}
    for i in range(0, len(tab)):
        try:
            tab_first_half = tab[i][[0, 1]]
            tab_second_half = tab[i][[2, 3]]
            tab_first_half.columns = ['attribute', 'value']
            tab_second_half.columns = ['attribute', 'value']
            output_dict[i] = pd.concat([tab_first_half, tab_second_half], ignore_index=True)
        except Exception:
            print('error in file {} page {}'.format(filename, i+1))
    return output_dict

def make_all_patients_location_table(DW_TABLES):
    """Create patient location summary table.
    
    Args:
        DW_TABLES (dict): Dictionary of patient word document tables
        
    Returns:
        pd.DataFrame: Table with patient location information
    """
    state_df = pd.DataFrame(columns=['patient_id', 'location_brain', 'location_body', 'gender'])
    
    for i in DW_TABLES:
        tmp_dw_table = DW_TABLES[i][0]
        tmp_dw_table = tmp_dw_table.set_index('attribute')

        if 'right' in tmp_dw_table.loc['ICH Location', 'value'].lower():
            location_brain = 'right'
            location_body = 'left'
        else:
            location_brain = 'left'
            location_body = 'right'
        gender = tmp_dw_table.loc['Sex', 'value'].lower() == 'male'

        state_df = state_df.append({
            'patient_id': i, 
            'location_brain': location_brain,
            'location_body': location_body, 
            'gender': gender
        }, ignore_index=True)

    state_df = state_df.set_index('patient_id')
    state_df['gender'] = state_df['gender'].astype(int)
    return state_df

def make_dw_df(DW_TABLES):
    """Create daily patient status DataFrame.
    
    Args:
        DW_TABLES (dict): Dictionary of patient word document tables
        
    Returns:
        pd.DataFrame: Table with daily patient status
    """
    dw_df = pd.DataFrame(columns=[
        'Patient_ID', 'Day_Num', 'location_brain',
        'age', 'sex', 'nihss', 'ventilation', 'Delirium_stat'
    ])

    coma_patients = 0

    for i in DW_TABLES:
        for j in DW_TABLES[i].keys():
            tmp_dw_table = DW_TABLES[i][j]
            tmp_dw_table = tmp_dw_table.set_index('attribute')

            # Process location and demographics
            location_brain = 'right' if 'right' in tmp_dw_table.loc['ICH Location', 'value'].lower() else 'left'
            gender = tmp_dw_table.loc['Sex', 'value'].lower() == 'male'
            age = int(tmp_dw_table.loc['Age', 'value'])
            nihss = int(tmp_dw_table.loc['Initial NIHSS (affected arm)', 'value'])
            
            # Process ventilation status
            ventilation = str(tmp_dw_table.loc['Mechanical ventilation', 'value'])
            ventilation = ventilation.replace(" ", "").lower()
            ventilation = "yes" in ventilation

            # Process delirium status
            try:
                delirium_stat_text = tmp_dw_table.loc['Delirium status', 'value'].lower()
                if "yes" in delirium_stat_text:
                    delirium_stat = True
                elif "no" in delirium_stat_text:
                    delirium_stat = False
                else:
                    coma_patients += 1
                    delirium_stat = None
            except:
                delirium_stat = None

            dw_df = dw_df.append({
                'Patient_ID': i, 
                'Day_Num': j+1, 
                'location_brain': location_brain,
                'age': age, 
                'sex': gender, 
                'nihss': nihss, 
                'ventilation': ventilation,
                'Delirium_stat': delirium_stat
            }, ignore_index=True)

    print("# of comatosed patients: ", coma_patients)

    dw_df = dw_df.set_index('Patient_ID')
    dw_df['sex'] = dw_df['sex'].astype(int)
    dw_df['ventilation'] = dw_df['ventilation'].astype(int)
    return dw_df

def preprocess_word_docs(data_path):
    """Read and process word documents.
    
    Args:
        data_path (str): Path to word documents
        
    Returns:
        tuple: (patient_location_df, dw_df) - Location and daily status DataFrames
    """
    # Read all DW tables
    dw_tables = {}  # key: patient ID, value: {key: day, value: table in doc for that day}
    for i in range(1, 31):
        dw_tables[i] = read_dw_word_files(data_path + 'DW'+str(i)+'.docx')

    # Process tables into DataFrames
    patient_location_df = make_all_patients_location_table(dw_tables)
    dw_df = make_dw_df(dw_tables)

    # Final processing
    dw_df = pd.get_dummies(dw_df, columns=["location_brain"], prefix=["ICH"])
    dw_df = dw_df.drop(columns=["ICH_left"])
    dw_df = dw_df.dropna(subset=["Delirium_stat"])
    dw_df['Delirium_stat'] = dw_df['Delirium_stat'].astype(int)

    return patient_location_df, dw_df

# =====================
# ICH Data Processing
# =====================

def propagate_value(df_col):
    """Replace nulls with column median.
    
    Args:
        df_col (pd.Series): Column to process
        
    Returns:
        pd.Series: Processed column with nulls filled
    """
    return df_col.fillna(value=df_col.median())

def create_ICH_data_dict(mapping_patient_study, ich_df):
    """Create dictionary of ICH data by patient.
    
    Args:
        mapping_patient_study (pd.DataFrame): Mapping between patient and study IDs
        ich_df (pd.DataFrame): ICH data
        
    Returns:
        dict: Dictionary of ICH data by patient
    """
    data_dict = {}
    columns_to_propagate = [
        'age', 'ich_score', 'nihss_total', 'nihss_arm_l', 'nihss_arm_r',
        'ich_max_volume', 'ivh', 'ich_side', 'sex'
    ]
    
    for patient_id in range(1, NUM_PATIENTS):
        study_id = mapping_patient_study.loc[patient_id].study_id
        data_dict[patient_id] = ich_df.loc[ich_df.study_id == study_id].copy(deep=True)
        data_dict[patient_id]['Patient_ID'] = int(patient_id)
        
        # Fill missing values
        for col in columns_to_propagate:
            data_dict[patient_id][col] = propagate_value(data_dict[patient_id][col])
            
        data_dict[patient_id] = data_dict[patient_id].drop(columns=['dw_id'])
    return data_dict

def make_patient_location_df(_data_dict):
    """
    Inputs:
        Dictionary of panda tables from all patients word documents
    Returns:
        A panda table where each row corresponds to one patient, and contains patient id, affected side of brain,
        affected side of the body, and patient's gender (0 stands for female, 1 stands for male)
    """
    state_df = pd.DataFrame(columns=['patient_id', 'location_brain', 'location_body', 'gender'])

    for i in _data_dict:
        patient_df = _data_dict[i]

        ich_side_unique = patient_df['ich_side'].unique()
        assert len(ich_side_unique == 1)
        ich_side = ich_side_unique[0]
        if ich_side == 1:
            location_brain = 'right'
            location_body = 'left'
        elif ich_side == 2:
            location_brain = 'left'
            location_body = 'right'
        else:
            raise ValueError(f"ich side {ich_side} not recognized")

        sex_unique = patient_df['sex'].unique()
        assert len(sex_unique == 1)
        sex = sex_unique[0]
        gender = int(sex) == 1

        state_df = state_df.append({'patient_id': i, 'location_brain': location_brain,
                                    'location_body': location_body, 'gender': gender}, ignore_index=True)

    state_df = state_df.set_index('patient_id')
    state_df['gender'] = state_df['gender'].astype(int)

    return state_df

def rename_affected_sides_ich(df, patient_location_df):
    """Rename arm columns to affected/non-affected.
    
    Args:
        df (pd.DataFrame): ICH patient data
        patient_location_df (pd.DataFrame): Patient location information
        
    Returns:
        pd.DataFrame: DataFrame with renamed columns
    """
    for i, r in df.iterrows():
        patient_id = r['Patient_ID']
        aff_side_left_bool = (patient_location_df.loc[patient_id]['location_body'] == 'left')

        df.loc[i, 'nihss_arm_affected'] = (
            df.loc[i, 'nihss_arm_l'] * aff_side_left_bool +
            df.loc[i, 'nihss_arm_r'] * (1-aff_side_left_bool)
        )
        df.loc[i, 'nihss_arm_not_affected'] = (
            df.loc[i, 'nihss_arm_l'] * (1-aff_side_left_bool) +
            df.loc[i, 'nihss_arm_r'] * aff_side_left_bool
        )

    df = df.drop(columns=['nihss_arm_l', 'nihss_arm_r'])
    return df

#######################
# Main Processing Functions
#######################

def preprocess_ich_data(data_path, patient_location_df=None):
    """
    Read the ICH data
    Input:
        data_path: path to ICH csv file
        patient_location_df: output of preprocess_word_docs that specifies location of ICH
    Output:
        ich_patient_clean: ICH data that has been cross-checked with dw_
    """
    ich_file = "ICHdb-5-16-21-dw-only.csv"
    ich_df = pd.read_csv(os.path.join(data_path, ich_file))
    ich_df['ich_volume'] = ich_df['ich_volume'].astype(float)
    ich_df['ich_volume_2'] = ich_df['ich_volume_2'].astype(float)
    ich_df['ich_volume_3'] = ich_df['ich_volume_3'].astype(float)

    # columns that Mike suggests us to use
    select_columns_alt = ['dw_id', 'study_id', 'delirium_expert', 'redcap_event_name', 'age', 'ich_score',
                          'ich_volume', 'ich_volume_2', 'ich_volume_3', 'nihss_total', 'nihss_arm_l', 'nihss_arm_r',
                          'ivh', 'ich_side', 'sex']

    ich_df = ich_df[select_columns_alt]
    ich_df['ich_max_volume'] = ich_df[['ich_volume', 'ich_volume_2', 'ich_volume_3']].max(axis=1)

    # create a mapping between the patient id (column 'dw_id') and the study id:
    mapping_patient_study = ich_df[['dw_id', 'study_id']].loc[ich_df.dw_id.notna()].set_index('dw_id')

    data_dict = create_ICH_data_dict(mapping_patient_study, ich_df)
    patient_location_df = make_patient_location_df(data_dict)

    ich_patient = pd.DataFrame()
    for patient_id in range(1, NUM_PATIENTS):
        df_non_null_status = data_dict[patient_id].loc[data_dict[patient_id].delirium_expert.notna()]
        df_non_null_status['Day_Num'] = df_non_null_status['redcap_event_name'].str.split(
            'hospital_day_(.*)_arm', expand=True)[1].dropna().astype(int)
        df_non_null_status['delirium_stat'] = df_non_null_status['delirium_expert']
        ich_patient = ich_patient.append(df_non_null_status, ignore_index=True)

    ich_patient = ich_patient.drop(columns=['ich_volume', 'ich_volume_2', 'ich_volume_3', 'delirium_expert', 
                                          'redcap_event_name', 'study_id'])

    ich_patient_clean = rename_affected_sides_ich(ich_patient, patient_location_df)
    ich_patient_clean = ich_patient.drop(columns=["nihss_arm_l", "nihss_arm_r"])

    ich_patient_clean = ich_patient_clean.dropna().reset_index(drop=True)
    populate_REFERENCE_DAYs(ich_patient_clean)
    return patient_location_df, ich_patient_clean

def preprocess_actigraph_data(data_path, patient_location_df, all_times=True,
                              removeOutliers=True, use_autoencoder=False):
    """
    Read and process actigraph data
    """
    ACTIGRAPH_FILES = load_actigraph_files(data_path)
    ACTIGRAPH_FILES = rename_columns_actigraph_files(ACTIGRAPH_FILES, patient_location_df)

    if removeOutliers:
        ACTIGRAPH_FILES = remove_outliers(ACTIGRAPH_FILES)

    ACTIGRAPH_FILES = get_day_patient(ACTIGRAPH_FILES)

    actigraph_features_df = pd.DataFrame(columns=['Patient_ID', 'date'])
    time_ranges = ['1_to_1', 'morning', 'afternoon', 'night'] if all_times else ['afternoon']

    for patient_id in tqdm(range(1, NUM_PATIENTS)):
        min_date = ACTIGRAPH_FILES[patient_id]['date'].min().date()
        populate_reference_dates(patient_id, min_date)
        control_date = get_control_day(patient_id)
        for day, date_i in enumerate(ACTIGRAPH_FILES[patient_id].date.dt.date.unique()):
            tmp = pd.Series()
            if use_autoencoder:
                tmp = tmp.append(autoencode(ACTIGRAPH_FILES, patient_id, date_i))
            for time_range in time_ranges:
                tmp = tmp.append(construct_feature(min_at_rest, time_range, ACTIGRAPH_FILES, patient_id, date_i))
                tmp = tmp.append(construct_feature(dynamic_time_warping_arms, time_range, ACTIGRAPH_FILES, patient_id, date_i))
                tmp = tmp.append(construct_feature_two_day(dynamic_time_warping_two_day, time_range, ACTIGRAPH_FILES, 
                                                         patient_id, date_i, control_date))

            tmp['Patient_ID'] = patient_id
            tmp['Day_Num'] = (date_i - min_date).days + 1
            tmp['date'] = date_i

            actigraph_features_df = actigraph_features_df.append(tmp, ignore_index=True)
    return actigraph_features_df

#######################
# Full Pipeline
#######################

def full_preprocess(data_path, output_path, include_actigraph=True,
                    all_times=True, removeOutliers=True, use_autoencoder=False):
    """
    Run complete preprocessing pipeline for both ICH and Actigraph data
    """
    print("preprocessing ICH data...")
    patient_location_df, ich_data = preprocess_ich_data(data_path)
    all_data = ich_data

    if include_actigraph:
        print("preprocessing actigraph data...")
        actigraph_features_df = preprocess_actigraph_data(data_path, patient_location_df, all_times, 
                                                        removeOutliers, use_autoencoder)
        # merge right so that only days with delirium measurement are included
        all_data = pd.merge(actigraph_features_df, all_data, on=['Patient_ID', 'Day_Num'], how='right')

    all_data = all_data[all_data['Day_Num'] != 1]  # drop first day b/c NaN
    all_data = all_data.dropna()
    all_data.to_csv(output_path, index=False)

    return all_data

#######################
# Main Entry Point
#######################

if __name__ == "__main__":
    path = "data/stroke_data/"
    full_preprocess(path, "data/test.csv", include_actigraph=False)

